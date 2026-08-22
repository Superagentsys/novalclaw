#!/usr/bin/env python3
"""
报告分段合并工具

将分段生成的 Markdown 报告片段按序号合并为完整的 Word 文档。
解决大模型上下文长度限制导致的长报告生成截断或质量下降问题。

工作流程:
  1. 审查过程中，按章节分段生成 Markdown 文件到 segments/ 目录
     （如 01_管理层摘要.md, 02_合同基本信息.md, ...）
  2. 所有分段生成完毕后，调用本脚本合并为最终 Word 报告

用法:
    python scripts/merge_report_segments.py <segments_dir> -o <output.docx>
    python scripts/merge_report_segments.py result/软件采购_2026-07-18/segments \
        -o "result/软件采购_2026-07-18/风险评估报告_软件采购_2026-07-18.docx" \
        -t "软件采购合同风险评估报告"

参数:
    segments_dir : 包含分段 Markdown 文件的目录
    -o/--output  : 输出 Word 文件路径
    -t/--title   : 报告标题（用于标题页，可选）
    --no-pagebreak : 分段之间不插入分页符
"""

import argparse
import re
import sys
from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _natural_key(path: Path) -> List:
    """自然排序键，使 01_、02_...10_ 正确排序"""
    return [int(text) if text.isdigit() else text.lower()
            for text in re.split(r'(\d+)', path.name)]


def _add_inline_runs(paragraph, text: str):
    """处理行内格式：**粗体**、`代码`、*斜体*"""
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _render_table(rows: List[str], doc: Document):
    """渲染 Markdown 表格为 Word 表格"""
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip('|').split('|')]
        parsed.append(cells)

    # 移除分隔符行（---|:---:|---）
    if len(parsed) >= 2 and all(re.match(r'^[-:]+$', c) for c in parsed[1] if c):
        parsed.pop(1)

    if not parsed or not parsed[0]:
        return

    n_cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=n_cols)
    table.style = 'Light Grid Accent 1'

    for i, row in enumerate(parsed):
        for j in range(n_cols):
            cell_text = row[j] if j < len(row) else ''
            cell = table.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            _add_inline_runs(p, cell_text)
            if i == 0:
                for run in p.runs:
                    run.bold = True


def parse_markdown_to_docx(md_text: str, doc: Document):
    """将 Markdown 文本解析并写入 Word 文档"""
    lines = md_text.split('\n')
    i = 0
    table_rows: List[str] = []
    in_code_block = False
    code_lines: List[str] = []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            _render_table(table_rows, doc)
            table_rows = []

    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith('```'):
            if in_code_block:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                flush_table()
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 表格行收集
        if line.strip().startswith('|') and line.strip().endswith('|'):
            table_rows.append(line.strip())
            i += 1
            continue
        else:
            flush_table()

        # 标题
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
        # 分隔线
        elif line.strip() == '---':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('—' * 30)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        # 无序列表
        elif re.match(r'^\s*[-*]\s+', line):
            text = re.sub(r'^\s*[-*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            _add_inline_runs(p, text)
        # 有序列表
        elif re.match(r'^\s*\d+\.\s+', line):
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            _add_inline_runs(p, text)
        # 引用
        elif line.strip().startswith('>'):
            text = line.strip()[1:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(20)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        # 空行
        elif line.strip() == '':
            pass
        # 普通段落
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, line.strip())

        i += 1

    # 处理剩余
    if in_code_block and code_lines:
        p = doc.add_paragraph()
        run = p.add_run('\n'.join(code_lines))
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    flush_table()


def merge_segments(segments_dir: str, output_path: str,
                   title: str = None, page_break: bool = True) -> int:
    """合并分段 Markdown 文件为 Word 文档"""
    seg_dir = Path(segments_dir)
    if not seg_dir.exists():
        print(f"错误: 分段目录不存在: {seg_dir}", file=sys.stderr)
        return 1

    segments = sorted(seg_dir.glob('*.md'), key=_natural_key)
    if not segments:
        print(f"错误: 分段目录中没有 .md 文件: {seg_dir}", file=sys.stderr)
        return 1

    doc = Document()

    # 默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)

    # 标题页
    if title:
        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    print(f"共发现 {len(segments)} 个分段文件，按顺序合并：")
    for idx, seg in enumerate(segments):
        print(f"  [{idx + 1}/{len(segments)}] {seg.name}")
        content = seg.read_text(encoding='utf-8')
        parse_markdown_to_docx(content, doc)
        if page_break and idx < len(segments) - 1:
            doc.add_page_break()

    # 确保输出目录存在
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc.save(output_path)
    print(f"\n✓ 合并完成: {output_path}")
    print(f"  共合并 {len(segments)} 个分段")
    return 0


def main():
    parser = argparse.ArgumentParser(
        description='合并分段 Markdown 报告为 Word 文档'
    )
    parser.add_argument('segments_dir', help='分段 Markdown 文件目录')
    parser.add_argument('-o', '--output', required=True,
                        help='输出 Word 文件路径')
    parser.add_argument('-t', '--title', help='报告标题（标题页）')
    parser.add_argument('--no-pagebreak', action='store_true',
                        help='分段之间不插入分页符')
    args = parser.parse_args()

    sys.exit(merge_segments(
        args.segments_dir, args.output,
        title=args.title, page_break=not args.no_pagebreak
    ))


if __name__ == '__main__':
    main()

#!/usr/bin/env python3
"""
合同版本对比工具

对比两个合同版本，识别新增、删除、修改的条款，生成差异报告。
"""

import sys
import argparse
import json
import re
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple


def extract_from_pdf(file_path: str) -> str:
    """从 PDF 文件中提取文本（pdfplumber 优先，PyPDF2 备用）"""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    text_parts.append(text)
        if text_parts:
            return "\n\n".join(text_parts)
    except ImportError:
        pass
    except Exception:
        pass

    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                text_parts.append(text)
        return "\n\n".join(text_parts)
    except Exception as e:
        raise Exception(f"PDF 文件读取失败: {str(e)}")


def extract_from_docx(file_path: str) -> str:
    """从 Docx 文件中提取文本（含表格内容）"""
    try:
        from docx import Document

        doc = Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            table_lines = []
            for i, row in enumerate(table.rows):
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    if i == 0:
                        table_lines.append("[表格] " + " | ".join(row_texts))
                    else:
                        table_lines.append(" | ".join(row_texts))
            if table_lines:
                text_parts.append("\n".join(table_lines))

        return "\n\n".join(text_parts)

    except Exception as e:
        raise Exception(f"Docx 文件读取失败: {str(e)}")


def extract_text(input_file: str) -> str:
    """
    从合同文件中提取文本内容

    参数:
        input_file: 合同文件路径，支持 PDF 或 Docx 格式

    返回:
        提取的纯文本内容
    """
    file_path = Path(input_file)

    if not file_path.exists():
        raise FileNotFoundError(f"文件不存在: {input_file}")

    if file_path.suffix.lower() == ".pdf":
        return extract_from_pdf(input_file)
    elif file_path.suffix.lower() in [".docx", ".doc"]:
        return extract_from_docx(input_file)
    else:
        raise ValueError(
            f"不支持的文件格式: {file_path.suffix}。仅支持 PDF 和 Docx 格式。"
        )


def split_paragraphs(text: str) -> List[str]:
    """
    将文本分割为段落

    参数:
        text: 文本内容

    返回:
        段落列表
    """
    # 按双换行符分割段落
    paragraphs = text.split('\n\n')

    # 过滤空段落，并去除首尾空白
    paragraphs = [p.strip() for p in paragraphs if p.strip()]

    return paragraphs


def normalize_text(text: str) -> str:
    """
    规范化文本（去除多余空格、换行等）

    参数:
        text: 原始文本

    返回:
        规范化后的文本
    """
    # 去除首尾空白
    text = text.strip()

    # 将多个连续空格替换为单个空格
    text = re.sub(r'\s+', ' ', text)

    return text


def compare_paragraphs(paragraphs1: List[str], paragraphs2: List[str]) -> List[Dict]:
    """
    对比两个版本的段落，识别差异

    参数:
        paragraphs1: 版本 1 的段落列表
        paragraphs2: 版本 2 的段落列表

    返回:
        差异列表
    """
    import difflib

    changes = []

    # 使用 difflib 的 SequenceMatcher 进行对比
    matcher = difflib.SequenceMatcher(None, paragraphs1, paragraphs2)

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'replace':
            # 替换：段落被修改
            for idx, (p1, p2) in enumerate(zip(paragraphs1[i1:i2], paragraphs2[j1:j2])):
                if normalize_text(p1) != normalize_text(p2):
                    changes.append({
                        'type': 'modified',
                        'paragraph_index': i1 + idx + 1,
                        'original_text': p1,
                        'new_text': p2
                    })
        elif tag == 'delete':
            # 删除：段落被移除
            for idx, p in enumerate(paragraphs1[i1:i2]):
                changes.append({
                    'type': 'removed',
                    'paragraph_index': i1 + idx + 1,
                    'text': p
                })
        elif tag == 'insert':
            # 新增：段落被添加
            for idx, p in enumerate(paragraphs2[j1:j2]):
                changes.append({
                    'type': 'added',
                    'paragraph_index': j1 + idx + 1,
                    'text': p
                })
        # tag == 'equal'：无变化，不记录

    return changes


def generate_html_report(
    version1_file: str,
    version2_file: str,
    changes: List[Dict],
    output_file: str
) -> None:
    """
    生成 HTML 格式的差异报告

    参数:
        version1_file: 版本 1 文件路径
        version2_file: 版本 2 文件路径
        changes: 差异列表
        output_file: 输出文件路径
    """
    # 统计差异
    stats = {
        'total': len(changes),
        'added': len([c for c in changes if c['type'] == 'added']),
        'removed': len([c for c in changes if c['type'] == 'removed']),
        'modified': len([c for c in changes if c['type'] == 'modified'])
    }

    # 生成 HTML
    html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>合同版本对比报告</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f5f5f5;
        }}
        .container {{
            background-color: white;
            padding: 30px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #333;
            border-bottom: 3px solid #4CAF50;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #666;
            margin-top: 30px;
        }}
        .version-info {{
            background-color: #f9f9f9;
            padding: 15px;
            border-radius: 4px;
            margin: 20px 0;
        }}
        .statistics {{
            display: flex;
            gap: 20px;
            margin: 20px 0;
        }}
        .stat-box {{
            flex: 1;
            padding: 15px;
            border-radius: 4px;
            text-align: center;
            color: white;
        }}
        .stat-total {{ background-color: #2196F3; }}
        .stat-added {{ background-color: #4CAF50; }}
        .stat-removed {{ background-color: #f44336; }}
        .stat-modified {{ background-color: #FF9800; }}
        .stat-number {{
            font-size: 2em;
            font-weight: bold;
        }}
        .change {{
            margin: 20px 0;
            padding: 15px;
            border-radius: 4px;
            border-left: 4px solid #ccc;
        }}
        .change.added {{
            background-color: #e8f5e9;
            border-left-color: #4CAF50;
        }}
        .change.removed {{
            background-color: #ffebee;
            border-left-color: #f44336;
        }}
        .change.modified {{
            background-color: #fff3e0;
            border-left-color: #FF9800;
        }}
        .change-title {{
            font-weight: bold;
            margin-bottom: 10px;
        }}
        .change-text {{
            background-color: white;
            padding: 10px;
            border-radius: 4px;
            white-space: pre-wrap;
            word-wrap: break-word;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>合同版本对比报告</h1>

        <div class="version-info">
            <h2>版本信息</h2>
            <p><strong>版本 1:</strong> {Path(version1_file).name}</p>
            <p><strong>版本 2:</strong> {Path(version2_file).name}</p>
            <p><strong>对比时间:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        </div>

        <h2>差异统计</h2>
        <div class="statistics">
            <div class="stat-box stat-total">
                <div class="stat-number">{stats['total']}</div>
                <div>总计修改</div>
            </div>
            <div class="stat-box stat-added">
                <div class="stat-number">{stats['added']}</div>
                <div>新增</div>
            </div>
            <div class="stat-box stat-removed">
                <div class="stat-number">{stats['removed']}</div>
                <div>删除</div>
            </div>
            <div class="stat-box stat-modified">
                <div class="stat-number">{stats['modified']}</div>
                <div>修改</div>
            </div>
        </div>

        <h2>详细差异</h2>
"""

    # 添加差异详情
    for change in changes:
        if change['type'] == 'added':
            html_content += f"""
        <div class="change added">
            <div class="change-title">[新增] 段落 {change['paragraph_index']}</div>
            <div class="change-text">{change['text']}</div>
        </div>
"""
        elif change['type'] == 'removed':
            html_content += f"""
        <div class="change removed">
            <div class="change-title">[删除] 段落 {change['paragraph_index']}</div>
            <div class="change-text">{change['text']}</div>
        </div>
"""
        elif change['type'] == 'modified':
            html_content += f"""
        <div class="change modified">
            <div class="change-title">[修改] 段落 {change['paragraph_index']}</div>
            <div style="margin-bottom: 10px;">
                <strong>原文:</strong>
                <div class="change-text">{change['original_text']}</div>
            </div>
            <div>
                <strong>新文:</strong>
                <div class="change-text">{change['new_text']}</div>
            </div>
        </div>
"""

    html_content += """
    </div>
</body>
</html>
"""

    # 写入文件
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(html_content)


def compare_contracts(
    version1_file: str,
    version2_file: str,
    json_output: str = None,
    html_output: str = None
) -> Dict:
    """
    对比两个合同版本

    参数:
        version1_file: 版本 1 文件路径
        version2_file: 版本 2 文件路径
        json_output: JSON 输出文件路径（可选）
        html_output: HTML 输出文件路径（可选）

    返回:
        对比结果字典
    """
    # 提取文本
    text1 = extract_text(version1_file)
    text2 = extract_text(version2_file)

    # 分割段落
    paragraphs1 = split_paragraphs(text1)
    paragraphs2 = split_paragraphs(text2)

    # 对比段落
    changes = compare_paragraphs(paragraphs1, paragraphs2)

    # 统计
    stats = {
        'total': len(changes),
        'added': len([c for c in changes if c['type'] == 'added']),
        'removed': len([c for c in changes if c['type'] == 'removed']),
        'modified': len([c for c in changes if c['type'] == 'modified'])
    }

    # 构建结果
    result = {
        'version1_file': str(Path(version1_file)),
        'version2_file': str(Path(version2_file)),
        'compare_time': datetime.now().isoformat(),
        'changes': changes,
        'statistics': stats
    }

    # 输出 JSON
    if json_output:
        with open(json_output, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)

    # 输出 HTML
    if html_output:
        generate_html_report(version1_file, version2_file, changes, html_output)

    return result


def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(
        description="对比两个合同版本，生成差异报告"
    )
    parser.add_argument(
        "version1",
        help="版本 1 文件路径（PDF 或 Docx 格式）"
    )
    parser.add_argument(
        "version2",
        help="版本 2 文件路径（PDF 或 Docx 格式）"
    )
    parser.add_argument(
        "--json-output",
        help="JSON 输出文件路径（可选）"
    )
    parser.add_argument(
        "--html-output",
        help="HTML 输出文件路径（可选）"
    )
    parser.add_argument(
        "--summary",
        action="store_true",
        help="显示对比汇总统计"
    )

    args = parser.parse_args()

    try:
        # 执行对比
        result = compare_contracts(
            args.version1,
            args.version2,
            args.json_output,
            args.html_output
        )

        # 显示汇总
        if args.summary or (not args.json_output and not args.html_output):
            stats = result['statistics']
            print(f"\n=== 版本对比汇总 ===", file=sys.stderr)
            print(f"版本 1: {Path(args.version1).name}", file=sys.stderr)
            print(f"版本 2: {Path(args.version2).name}", file=sys.stderr)
            print(f"\n总计修改: {stats['total']} 处", file=sys.stderr)
            print(f"新增: {stats['added']} 处", file=sys.stderr)
            print(f"删除: {stats['removed']} 处", file=sys.stderr)
            print(f"修改: {stats['modified']} 处", file=sys.stderr)

            if args.json_output:
                print(f"\nJSON 报告已保存到: {args.json_output}", file=sys.stderr)
            if args.html_output:
                print(f"HTML 报告已保存到: {args.html_output}", file=sys.stderr)

        return 0

    except Exception as e:
        print(f"错误: {str(e)}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())

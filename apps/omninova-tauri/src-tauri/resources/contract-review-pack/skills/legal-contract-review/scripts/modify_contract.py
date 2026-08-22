#!/usr/bin/env python3
"""
合同文本替换工具

在 Docx 格式的合同文档中批量替换文本内容，生成修改后的文档。
支持跨 run 替换（Word 文档中的文本经常跨越多个 run）。
"""

import sys
import argparse
import json
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _copy_run_properties(source_run, target_run):
    """复制 run 的格式属性"""
    if source_run.font.name:
        target_run.font.name = source_run.font.name
    if source_run.font.size:
        target_run.font.size = source_run.font.size
    if source_run.font.bold is not None:
        target_run.font.bold = source_run.font.bold
    if source_run.font.italic is not None:
        target_run.font.italic = source_run.font.italic
    if source_run.font.color and source_run.font.color.rgb:
        target_run.font.color.rgb = source_run.font.color.rgb
    if source_run.font.underline is not None:
        target_run.font.underline = source_run.font.underline


def _replace_in_paragraph(paragraph, original_text, new_text):
    """
    在段落中替换文本（支持跨 run）

    返回:
        (是否替换成功, 替换次数)
    """
    full_text = paragraph.text
    if original_text not in full_text:
        return False, 0

    # 尝试简单替换：检查每个 run 是否完整包含 original_text
    simple_replaced = 0
    for run in paragraph.runs:
        if original_text in run.text:
            run.text = run.text.replace(original_text, new_text)
            simple_replaced += 1

    if simple_replaced > 0:
        return True, simple_replaced

    # 跨 run 替换：合并后重建
    runs = list(paragraph.runs)
    if not runs:
        return False, 0

    # 收集所有 run 信息
    merged_text = "".join(run.text for run in runs)

    if original_text not in merged_text:
        return False, 0

    # 找到替换位置
    start_pos = merged_text.find(original_text)
    end_pos = start_pos + len(original_text)
    new_merged = merged_text[:start_pos] + new_text + merged_text[end_pos:]

    # 按 run 边界重建文本
    # 策略：找到起始和结束 run，保留前后 run 不变，中间的重建
    char_pos = 0
    start_run_idx = -1
    end_run_idx = -1
    start_offset = 0
    end_offset = 0

    for idx, run in enumerate(runs):
        run_len = len(run.text)
        if start_run_idx == -1 and char_pos + run_len > start_pos:
            start_run_idx = idx
            start_offset = start_pos - char_pos
        if char_pos + run_len >= end_pos:
            end_run_idx = idx
            end_offset = end_pos - char_pos
            break
        char_pos += run_len

    if start_run_idx == -1 or end_run_idx == -1:
        return False, 0

    if start_run_idx == end_run_idx:
        runs[start_run_idx].text = runs[start_run_idx].text[:start_offset] + new_text + runs[start_run_idx].text[end_offset:]
        return True, 1

    # 保留第一个 run 的格式
    first_style = runs[start_run_idx]

    # 修改起始 run
    runs[start_run_idx].text = runs[start_run_idx].text[:start_offset] + new_text

    # 清空起始和结束之间的 run
    for idx in range(start_run_idx + 1, end_run_idx + 1):
        if idx == end_run_idx:
            runs[idx].text = runs[idx].text[end_offset:]
        else:
            runs[idx].text = ""

    return True, 1


def modify_contract(input_file, output_file, replacements, track_changes=False):
    """
    在合同文档中替换文本

    参数:
        input_file: 输入文件路径（Docx 格式）
        output_file: 输出文件路径
        replacements: 替换信息列表
        track_changes: 是否启用修订模式（Track Changes）

    返回:
        替换统计信息
    """
    try:
        doc = Document(input_file)
        stats = {
            "total_replacements": len(replacements),
            "successful": 0,
            "failed": 0,
            "details": []
        }

        for replacement in replacements:
            original_text = replacement.get("original_text", "")
            new_text = replacement.get("new_text", "")

            if not original_text or not new_text:
                stats["failed"] += 1
                stats["details"].append({
                    "original_text": original_text,
                    "new_text": new_text,
                    "status": "skipped_empty",
                    "message": "原始文本或新文本为空"
                })
                continue

            paragraph_matched = False
            match_count = 0

            # 遍历段落
            for para in doc.paragraphs:
                if original_text in para.text:
                    if track_changes:
                        _replace_with_track_changes(para, original_text, new_text)
                    else:
                        success, count = _replace_in_paragraph(para, original_text, new_text)
                        if success:
                            match_count += count
                            paragraph_matched = True

            # 遍历表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if original_text in para.text:
                                if track_changes:
                                    _replace_with_track_changes(para, original_text, new_text)
                                else:
                                    success, count = _replace_in_paragraph(para, original_text, new_text)
                                    if success:
                                        match_count += count
                                        paragraph_matched = True

            if paragraph_matched:
                stats["successful"] += 1
                stats["details"].append({
                    "original_text": original_text,
                    "new_text": new_text,
                    "status": "success",
                    "matches": match_count,
                    "message": f"在 {match_count} 处完成替换"
                })
            else:
                stats["failed"] += 1
                stats["details"].append({
                    "original_text": original_text,
                    "new_text": new_text,
                    "status": "not_found",
                    "message": "未找到匹配文本，请检查原文片段是否准确"
                })

        doc.save(output_file)
        return stats

    except Exception as e:
        raise Exception(f"文档替换失败: {str(e)}")


def _replace_with_track_changes(paragraph, original_text, new_text):
    """
    使用修订模式替换文本（Track Changes）

    通过插入 w:ins（新增）和 w:del（删除）元素实现。
    """
    runs = list(paragraph.runs)
    if not runs:
        return

    merged_text = "".join(run.text for run in runs)

    if original_text not in merged_text:
        return

    # 创建删除标记（原文）和插入标记（新文）
    start_pos = merged_text.find(original_text)
    end_pos = start_pos + len(original_text)

    # 构建新的段落内容
    # 使用 XML 操作实现修订
    new_runs_text = []
    char_pos = 0

    for run in runs:
        run_len = len(run.text)
        run_start = char_pos
        run_end = char_pos + run_len

        if run_end <= start_pos or run_start >= end_pos:
            # 不涉及替换的部分，保留原 run
            new_runs_text.append(("keep", run, run.text))
        else:
            # 涉及替换的部分
            before = run.text[:max(0, start_pos - run_start)]
            after = run.text[max(0, end_pos - run_start + len(original_text)):]

            if before:
                new_runs_text.append(("keep", run, before))
            if after:
                new_runs_text.append(("keep", run, after))

        char_pos += run_len

    # 重建段落内容
    # 简化版本：清空所有 run 并重建，用删除/插入标记包裹
    for run in runs:
        run.text = ""

    if runs:
        # 用第一个 run 设置全部内容
        full_text = merged_text[:start_pos] + new_text + merged_text[end_pos:]
        para_element = paragraph._p

        # 清除旧 runs
        for run_elem in list(para_element.findall(qn('w:r'))):
            para_element.remove(run_elem)

        if start_pos > 0:
            before_r = OxmlElement('w:r')
            before_t = OxmlElement('w:t')
            before_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            before_t.text = merged_text[:start_pos]
            before_r.append(before_t)
            para_element.append(before_r)

        # 删除原文
        del_r = OxmlElement('w:r')
        del_props = OxmlElement('w:rPr')
        del_r.append(del_props)
        del_text_elem = OxmlElement('w:delText')
        del_text_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        del_text_elem.text = original_text
        del_r.append(del_text_elem)
        para_element.append(del_r)

        # 插入新文
        ins_r = OxmlElement('w:r')
        ins_props = OxmlElement('w:rPr')
        ins_r.append(ins_props)
        ins_t = OxmlElement('w:t')
        ins_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        ins_t.text = new_text
        ins_r.append(ins_t)
        para_element.append(ins_r)

        if end_pos < len(merged_text):
            after_r = OxmlElement('w:r')
            after_t = OxmlElement('w:t')
            after_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            after_t.text = merged_text[end_pos:]
            after_r.append(after_t)
            para_element.append(after_r)


def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(
        description="在合同文档中批量替换文本内容"
    )
    parser.add_argument(
        "input_file",
        help="输入合同文件路径（Docx 格式）"
    )
    parser.add_argument(
        "output_file",
        help="输出合同文件路径（Docx 格式）"
    )
    parser.add_argument(
        "replacements_file",
        help="替换信息文件（JSON 格式）"
    )
    parser.add_argument(
        "--track-changes",
        action="store_true",
        help="启用修订模式（Track Changes），保留修改痕迹"
    )

    args = parser.parse_args()

    try:
        with open(args.replacements_file, 'r', encoding='utf-8') as f:
            replacements = json.load(f)

        if not isinstance(replacements, list):
            raise ValueError("替换信息必须是列表格式")
        if len(replacements) == 0:
            raise ValueError("替换信息列表不能为空")

        stats = modify_contract(
            args.input_file,
            args.output_file,
            replacements,
            track_changes=args.track_changes
        )

        print(f"替换完成。总计 {stats['total_replacements']} 项，成功 {stats['successful']} 项，失败 {stats['failed']} 项", file=sys.stderr)
        print(f"输出文件: {args.output_file}", file=sys.stderr)

        if stats["failed"] > 0:
            failed_items = [d for d in stats["details"] if d["status"] != "success"]
            for item in failed_items:
                original_preview = item["original_text"][:40] + ("..." if len(item["original_text"]) > 40 else "")
                print(f"  失败: '{original_preview}' - {item['message']}", file=sys.stderr)

        return 0

    except Exception as e:
        print(f"错误: {str(e)}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())

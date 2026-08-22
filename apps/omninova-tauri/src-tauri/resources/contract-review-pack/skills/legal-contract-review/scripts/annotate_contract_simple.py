#!/usr/bin/env python3
"""
合同风险标注工具（简化版）

在 Docx 合同原文中高亮标注风险条款关键词。
- 高风险：黄色高亮
- 中风险：绿色高亮
- 低风险：灰色高亮

用法:
    python scripts/annotate_contract_simple.py <input.docx> <output.docx> <annotations.json>

annotations.json 格式:
    [
      {"keyword": "匹配关键词", "risk_level": "high|medium|low", "comment": "说明"},
      ...
    ]
"""

import sys
import json
import argparse
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX


RISK_HIGHLIGHT = {
    "high": WD_COLOR_INDEX.YELLOW,
    "medium": WD_COLOR_INDEX.BRIGHT_GREEN,
    "low": WD_COLOR_INDEX.GRAY_25,
}

RISK_LABEL = {
    "high": "[高风险]",
    "medium": "[中风险]",
    "low": "[低风险]",
}


def highlight_in_paragraph(paragraph, keyword, risk_level, comment):
    """在段落中查找关键词并高亮，返回标注数量"""
    full_text = paragraph.text
    if keyword not in full_text:
        return 0

    count = 0
    # 重建段落 runs
    # 策略：将段落文本按关键词分割，重建 runs
    parts = full_text.split(keyword)
    if len(parts) <= 1:
        return 0

    # 清空原段落
    for run in paragraph.runs:
        run.text = ""

    # 重建
    highlight_color = RISK_HIGHLIGHT.get(risk_level, WD_COLOR_INDEX.YELLOW)
    label = RISK_LABEL.get(risk_level, "")

    for i, part in enumerate(parts):
        if part:
            run = paragraph.add_run(part)
        if i < len(parts) - 1:
            # 添加高亮的关键词
            hl_run = paragraph.add_run(keyword)
            hl_run.font.highlight_color = highlight_color
            # 添加标注说明
            if comment:
                note_run = paragraph.add_run(f" {label}{comment} ")
                note_run.font.highlight_color = highlight_color
                note_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00) if risk_level == "high" else RGBColor(0x00, 0x60, 0x00)
            count += 1

    return count


def annotate_contract(input_path, output_path, annotations):
    """标注合同"""
    doc = Document(input_path)

    total_count = 0
    for ann in annotations:
        keyword = ann.get("keyword", "")
        risk_level = ann.get("risk_level", "medium")
        comment = ann.get("comment", "")

        if not keyword:
            continue

        for paragraph in doc.paragraphs:
            total_count += highlight_in_paragraph(paragraph, keyword, risk_level, comment)

        # 表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        total_count += highlight_in_paragraph(paragraph, keyword, risk_level, comment)

    doc.save(output_path)
    return total_count


def main():
    parser = argparse.ArgumentParser(description="合同风险标注工具（简化版）")
    parser.add_argument("input", help="输入合同 Docx 文件")
    parser.add_argument("output", help="输出标注后 Docx 文件")
    parser.add_argument("annotations", help="标注信息 JSON 文件")
    args = parser.parse_args()

    with open(args.annotations, "r", encoding="utf-8") as f:
        annotations = json.load(f)

    print(f"加载 {len(annotations)} 条标注信息")
    count = annotate_contract(args.input, args.output, annotations)
    print(f"✓ 标注完成: {args.output}")
    print(f"  共标注 {count} 处风险条款")


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""
合同文本提取工具

支持从 PDF 和 Docx 格式的合同文件中提取纯文本内容。
PDF 提取优先级：pdfplumber > PyPDF2 > OCR（Tesseract）
"""

import sys
import argparse
import warnings
from pathlib import Path


def _extract_with_pdfplumber(file_path: str) -> str:
    """使用 pdfplumber 提取 PDF 文本（优选引擎）"""
    import pdfplumber
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text and text.strip():
                text_parts.append(text)
    if not text_parts:
        raise Exception("pdfplumber 未提取到任何文本")
    return "\n\n".join(text_parts)


def _extract_with_pypdf2(file_path: str) -> str:
    """使用 PyPDF2 提取 PDF 文本（备用引擎）"""
    from PyPDF2 import PdfReader
    reader = PdfReader(file_path)
    text_parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text and text.strip():
            text_parts.append(text)
    if not text_parts:
        raise Exception("PyPDF2 未提取到任何文本，文件可能是扫描件")
    return "\n\n".join(text_parts)


def _extract_with_ocr(file_path: str) -> str:
    """使用 Tesseract OCR 提取扫描件文本（最后降级方案）"""
    try:
        import pytesseract
        from pdf2image import convert_from_path

        images = convert_from_path(file_path, first_page=1, last_page=50)
        text_parts = []
        for i, image in enumerate(images):
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')
            if text.strip():
                text_parts.append(text)

        if not text_parts:
            raise Exception("OCR 未识别到任何文本")
        return "\n\n".join(text_parts)

    except ImportError as e:
        raise Exception(
            f"扫描件 PDF 文本提取失败。请确保已安装 OCR 依赖："
            f"pip install pytesseract pdf2image，并安装 Tesseract OCR 及中文语言包。"
            f"原始错误: {str(e)}"
        )


def extract_from_pdf(file_path: str) -> str:
    """
    从 PDF 文件中提取文本

    提取策略（按优先级降级）：
    1. pdfplumber（首选，对中文和表格支持更好）
    2. PyPDF2（备用，兼容性广泛）
    3. Tesseract OCR（扫描件最终方案）
    """
    engines = [
        ("pdfplumber", _extract_with_pdfplumber),
        ("PyPDF2", _extract_with_pypdf2),
        ("OCR", _extract_with_ocr),
    ]

    last_error = None
    for engine_name, engine_func in engines:
        try:
            return engine_func(file_path)
        except ImportError as e:
            last_error = str(e)
            continue
        except Exception as e:
            last_error = str(e)
            if engine_name == "OCR":
                raise Exception(
                    f"PDF 文本提取失败，所有引擎均无法提取文本。"
                    f"建议检查文件是否完整、是否为加密 PDF 或图片分辨率是否过低。"
                    f"最后错误: {last_error}"
                )
            continue

    raise Exception(f"PDF 文件读取失败: {last_error}")


def extract_from_docx(file_path: str) -> str:
    """从 Docx 文件中提取文本（含表格内容）"""
    try:
        from docx import Document

        doc = Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            table_lines = []
            for i, row in enumerate(table.rows):
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    if i == 0:
                        table_lines.append("[表格] " + " | ".join(row_texts))
                    else:
                        table_lines.append(" | ".join(row_texts))
            if table_lines:
                text_parts.append("\n".join(table_lines))

        return "\n\n".join(text_parts)

    except Exception as e:
        raise Exception(f"Docx 文件读取失败: {str(e)}")


def extract_text(input_file: str) -> str:
    """
    从合同文件中提取文本内容

    参数:
        input_file: 合同文件路径，支持 PDF 或 Docx 格式

    返回:
        提取的纯文本内容

    异常:
        ValueError: 不支持的文件格式
        Exception: 文件读取失败
    """
    file_path = Path(input_file)

    if not file_path.exists():
        raise FileNotFoundError(f"文件不存在: {input_file}")

    file_extension = file_path.suffix.lower()

    if file_extension == ".pdf":
        return extract_from_pdf(input_file)
    elif file_extension in [".docx", ".doc"]:
        return extract_from_docx(input_file)
    else:
        raise ValueError(
            f"不支持的文件格式: {file_extension}。仅支持 PDF 和 Docx 格式。"
        )


def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(
        description="从合同文件（PDF/Docx）中提取文本内容"
    )
    parser.add_argument(
        "input_file",
        help="合同文件路径，支持 PDF 或 Docx 格式"
    )
    parser.add_argument(
        "-o", "--output",
        help="输出文件路径（可选），不指定则打印到标准输出"
    )

    args = parser.parse_args()

    try:
        text = extract_text(args.input_file)

        if args.output:
            with open(args.output, "w", encoding="utf-8") as f:
                f.write(text)
            print(f"文本已提取并保存到: {args.output}", file=sys.stderr)
        else:
            print(text)

        return 0

    except Exception as e:
        print(f"错误: {str(e)}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())

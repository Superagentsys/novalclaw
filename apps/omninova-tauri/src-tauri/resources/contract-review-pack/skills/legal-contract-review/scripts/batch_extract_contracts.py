#!/usr/bin/env python3
"""
批量合同文本提取工具

批量从多个 PDF/Docx 格式的合同文件中提取文本内容，生成汇总报告。
"""

import sys
import argparse
import json
import time
from pathlib import Path
from typing import List, Dict
from concurrent.futures import ThreadPoolExecutor, as_completed


def extract_from_pdf(file_path: str) -> str:
    """从 PDF 文件中提取文本（pdfplumber 优先，PyPDF2 备用）"""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    text_parts.append(text)
        if text_parts:
            return "\n\n".join(text_parts)
    except ImportError:
        pass
    except Exception:
        pass

    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                text_parts.append(text)
        return "\n\n".join(text_parts)
    except Exception as e:
        raise Exception(f"PDF 文件读取失败: {str(e)}")


def extract_from_docx(file_path: str) -> str:
    """从 Docx 文件中提取文本（含表格内容）"""
    try:
        from docx import Document

        doc = Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            table_lines = []
            for i, row in enumerate(table.rows):
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    if i == 0:
                        table_lines.append("[表格] " + " | ".join(row_texts))
                    else:
                        table_lines.append(" | ".join(row_texts))
            if table_lines:
                text_parts.append("\n".join(table_lines))

        return "\n\n".join(text_parts)

    except Exception as e:
        raise Exception(f"Docx 文件读取失败: {str(e)}")


def extract_text(input_file: str) -> Dict[str, str]:
    """
    从合同文件中提取文本内容

    参数:
        input_file: 合同文件路径，支持 PDF 或 Docx 格式

    返回:
        提取结果字典：
        {
            "file_name": 文件名,
            "file_path": 文件路径,
            "file_type": 文件类型 (pdf/docx),
            "text": 提取的文本内容,
            "char_count": 字符数,
            "status": "success/error",
            "error": 错误信息（如果有）
        }
    """
    file_path = Path(input_file)
    result = {
        "file_name": file_path.name,
        "file_path": str(file_path),
        "file_type": file_path.suffix.lower(),
        "text": "",
        "char_count": 0,
        "status": "success",
        "error": ""
    }

    try:
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {input_file}")

        if file_path.suffix.lower() == ".pdf":
            result["text"] = extract_from_pdf(input_file)
            result["file_type"] = "pdf"
        elif file_path.suffix.lower() in [".docx", ".doc"]:
            result["text"] = extract_from_docx(input_file)
            result["file_type"] = "docx"
        else:
            raise ValueError(
                f"不支持的文件格式: {file_path.suffix}。仅支持 PDF 和 Docx 格式。"
            )

        result["char_count"] = len(result["text"])
        result["status"] = "success"

    except Exception as e:
        result["status"] = "error"
        result["error"] = str(e)

    return result


def batch_extract(input_files: List[str], output_file: str = None, workers: int = 1) -> List[Dict]:
    """
    批量提取合同文本（支持并行处理）

    参数:
        input_files: 输入文件路径列表
        output_file: 输出文件路径（JSON 格式），可选
        workers: 并行工作线程数（默认为1，即串行处理）

    返回:
        提取结果列表
    """
    results = [None] * len(input_files)

    if workers <= 1:
        for i, input_file in enumerate(input_files):
            result = extract_text(input_file)
            results[i] = result
            print(f"{'✓' if result['status'] == 'success' else '✗'} {result['file_name']}: "
                  f"{result['status']} ({result['char_count']} 字符)", file=sys.stderr)
    else:
        print(f"使用 {workers} 个并行线程处理 {len(input_files)} 个文件...", file=sys.stderr)
        with ThreadPoolExecutor(max_workers=workers) as executor:
            future_to_index = {
                executor.submit(extract_text, f): i
                for i, f in enumerate(input_files)
            }
            for future in as_completed(future_to_index):
                idx = future_to_index[future]
                try:
                    result = future.result()
                    results[idx] = result
                    print(f"{'✓' if result['status'] == 'success' else '✗'} {result['file_name']}: "
                          f"{result['status']} ({result['char_count']} 字符)", file=sys.stderr)
                except Exception as e:
                    results[idx] = {
                        "file_name": Path(input_files[idx]).name,
                        "file_path": input_files[idx],
                        "status": "error",
                        "error": str(e)
                    }
                    print(f"✗ {Path(input_files[idx]).name}: error ({str(e)})", file=sys.stderr)

    if output_file:
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(results, f, ensure_ascii=False, indent=2)
        print(f"\n提取结果已保存到: {output_file}", file=sys.stderr)

    return results


def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(
        description="批量从合同文件（PDF/Docx）中提取文本内容"
    )
    parser.add_argument(
        "input_files",
        nargs='+',
        help="合同文件路径列表，支持 PDF 或 Docx 格式"
    )
    parser.add_argument(
        "-o", "--output",
        help="输出文件路径（JSON 格式），可选"
    )
    parser.add_argument(
        "--summary",
        action="store_true",
        help="显示提取汇总统计"
    )
    parser.add_argument(
        "-w", "--workers",
        type=int,
        default=1,
        help="并行工作线程数（默认为1），建议不超过CPU核心数"
    )

    args = parser.parse_args()

    try:
        start_time = time.time()
        results = batch_extract(args.input_files, args.output, workers=args.workers)

        if args.summary:
            success_count = sum(1 for r in results if r['status'] == 'success')
            error_count = sum(1 for r in results if r['status'] == 'error')
            total_chars = sum(r['char_count'] for r in results)

            print(f"\n=== 提取汇总 ===", file=sys.stderr)
            print(f"总文件数: {len(results)}", file=sys.stderr)
            print(f"成功: {success_count}", file=sys.stderr)
            print(f"失败: {error_count}", file=sys.stderr)
            print(f"总字符数: {total_chars}", file=sys.stderr)
            elapsed = time.time() - start_time
            print(f"耗时: {elapsed:.1f} 秒 (并行线程数: {args.workers})", file=sys.stderr)

        return 0

    except Exception as e:
        print(f"错误: {str(e)}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())

#!/usr/bin/env python3
"""
合同完整性检查工具

检查合同是否包含必备条款，按合同类型分类检查，输出缺失条款清单和完整度评分。
"""

import sys
import argparse
import json
import re
from pathlib import Path
from typing import Dict, List, Optional


def extract_from_pdf(file_path: str) -> str:
    """从 PDF 文件中提取文本"""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                text_parts.append(text)
        return "\n\n".join(text_parts)
    except Exception as e:
        raise Exception(f"PDF 文件读取失败: {str(e)}")


def extract_from_docx(file_path: str) -> str:
    """从 Docx 文件中提取文本（含表格）"""
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    text_parts.append(" | ".join(row_texts))

        return "\n\n".join(text_parts)
    except Exception as e:
        raise Exception(f"Docx 文件读取失败: {str(e)}")


def extract_text(input_file: str) -> str:
    """从合同文件中提取文本内容"""
    file_path = Path(input_file)
    if not file_path.exists():
        raise FileNotFoundError(f"文件不存在: {input_file}")
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return extract_from_pdf(input_file)
    elif ext in [".docx", ".doc"]:
        return extract_from_docx(input_file)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


CLAUSE_CHECKLIST = {
    "universal": [
        {
            "id": "parties",
            "name": "合同主体信息",
            "importance": "critical",
            "keywords": ["甲方", "乙方", "住所", "法定代表人", "统一社会信用代码"],
        },
        {
            "id": "subject_matter",
            "name": "合同标的/合作内容",
            "importance": "critical",
            "keywords": ["标的", "内容", "范围", "服务内容", "产品", "项目"],
        },
        {
            "id": "amount_price",
            "name": "价款与支付条款",
            "importance": "critical",
            "keywords": ["金额", "价款", "价格", "费用", "付款", "支付方式", "人民币"],
        },
        {
            "id": "performance",
            "name": "履行期限与方式",
            "importance": "critical",
            "keywords": ["期限", "交付", "履行", "完成时间", "工期", "进度"],
        },
        {
            "id": "breach",
            "name": "违约责任",
            "importance": "critical",
            "keywords": ["违约", "违约金", "赔偿", "违约方"],
        },
        {
            "id": "dispute",
            "name": "争议解决",
            "importance": "critical",
            "keywords": ["争议", "仲裁", "诉讼", "管辖", "法院"],
        },
        {
            "id": "force_majeure",
            "name": "不可抗力",
            "importance": "high",
            "keywords": ["不可抗力", "自然灾害", "政府行为", "战争"],
        },
        {
            "id": "notice",
            "name": "通知送达条款",
            "importance": "high",
            "keywords": ["通知", "送达", "地址", "联系人", "联系方式", "书面"],
        },
        {
            "id": "effective_term",
            "name": "合同生效与期限",
            "importance": "high",
            "keywords": ["生效", "有效期", "终止", "签订", "签署"],
        },
        {
            "id": "signature",
            "name": "签署盖章",
            "importance": "high",
            "keywords": ["签字", "盖章", "签章", "授权代表", "公章"],
        },
        {
            "id": "modification",
            "name": "合同变更条款",
            "importance": "medium",
            "keywords": ["变更", "修改", "补充", "修订"],
        },
        {
            "id": "severability",
            "name": "可分割性条款",
            "importance": "medium",
            "keywords": ["可分割", "无效", "部分无效"],
        },
    ],

    "nda": [
        {
            "id": "confidential_info",
            "name": "保密信息定义",
            "importance": "critical",
            "keywords": ["保密信息", "商业秘密", "机密", "定义", "包括"],
        },
        {
            "id": "confidential_period",
            "name": "保密期限",
            "importance": "critical",
            "keywords": ["保密期限", "保密期", "年", "终止后"],
        },
        {
            "id": "confidential_scope",
            "name": "保密义务范围",
            "importance": "critical",
            "keywords": ["保密义务", "不得", "披露", "泄露", "使用"],
        },
        {
            "id": "confidential_exemptions",
            "name": "保密豁免条款",
            "importance": "high",
            "keywords": ["除外", "豁免", "已公开", "第三方", "法律要求"],
        },
    ],

    "service": [
        {
            "id": "service_scope",
            "name": "服务范围与标准",
            "importance": "critical",
            "keywords": ["服务范围", "服务标准", "技术规范", "需求"],
        },
        {
            "id": "delivery",
            "name": "交付与验收条款",
            "importance": "critical",
            "keywords": ["交付", "验收", "确认", "签收", "测试"],
        },
        {
            "id": "sla",
            "name": "服务水平协议(SLA)",
            "importance": "high",
            "keywords": ["服务水平", "SLA", "可用性", "响应时间", "故障"],
        },
        {
            "id": "quality",
            "name": "质量保证条款",
            "importance": "high",
            "keywords": ["质量", "保证", "质保", "缺陷", "瑕疵"],
        },
        {
            "id": "ipr",
            "name": "知识产权归属",
            "importance": "high",
            "keywords": ["知识产权", "著作权", "专利权", "商标", "源代码"],
        },
    ],

    "procurement": [
        {
            "id": "goods_spec",
            "name": "货物规格与质量标准",
            "importance": "critical",
            "keywords": ["规格", "质量标准", "技术参数", "品牌", "型号"],
        },
        {
            "id": "delivery_schedule",
            "name": "交付时间与地点",
            "importance": "critical",
            "keywords": ["交货", "发货", "运输", "地点", "时间"],
        },
        {
            "id": "warranty",
            "name": "质保期与售后服务",
            "importance": "high",
            "keywords": ["质保", "保修", "售后", "维修", "退换"],
        },
        {
            "id": "inspection",
            "name": "验收与检验条款",
            "importance": "high",
            "keywords": ["验收", "检验", "检测", "合格", "不合格"],
        },
    ],

    "cooperation": [
        {
            "id": "cooperation_scope",
            "name": "合作范围与模式",
            "importance": "critical",
            "keywords": ["合作", "模式", "分工", "责任", "权益"],
        },
        {
            "id": "investment",
            "name": "投入与资源条款",
            "importance": "critical",
            "keywords": ["投入", "资源", "资金", "人员", "设备"],
        },
        {
            "id": "profit_sharing",
            "name": "收益分配条款",
            "importance": "critical",
            "keywords": ["收益", "利润", "分配", "分成", "结算"],
        },
        {
            "id": "ipr_cooperation",
            "name": "知识产权归属与使用",
            "importance": "critical",
            "keywords": ["知识产权", "开发", "共有", "许可", "授权"],
        },
        {
            "id": "exit",
            "name": "退出与清算条款",
            "importance": "high",
            "keywords": ["退出", "解散", "清算", "转让", "回购"],
        },
    ],
}


def check_clause_presence(text: str, keywords: List[str]) -> bool:
    """检查条款是否存在于文本中（多关键词综合判断）"""
    text_lower = text.lower()
    match_count = 0
    required_matches = max(1, len(keywords) // 2)

    for kw in keywords:
        if kw.lower() in text_lower:
            match_count += 1

    return match_count >= required_matches


def detect_contract_type(text: str) -> str:
    """自动检测合同类型"""
    type_keywords = {
        "nda": ["保密协议", "nda"],
        "service": ["服务合同", "服务协议", "技术服务", "软件开发"],
        "procurement": ["采购合同", "采购协议", "供货合同", "买卖合同"],
        "cooperation": ["合作合同", "合作协议", "战略合作", "联合开发"],
        "lease": ["租赁合同", "租赁协议", "出租"],
        "labor": ["劳动合同", "劳动协议", "劳务"],
    }

    text_head = text[:5000].lower()
    scores = {}
    for ctype, keywords in type_keywords.items():
        scores[ctype] = sum(1 for kw in keywords if kw.lower() in text_head)

    if max(scores.values()) == 0:
        return "general"
    return max(scores, key=scores.get)


def check_completeness(
    input_file: str,
    contract_type: Optional[str] = None
) -> Dict[str, any]:
    """
    检查合同完整性

    参数:
        input_file: 合同文件路径
        contract_type: 合同类型（可选，不指定则自动检测）

    返回:
        完整性检查结果字典
    """
    file_path = Path(input_file)
    text = extract_text(input_file)

    if contract_type is None:
        contract_type = detect_contract_type(text)

    clause_list = CLAUSE_CHECKLIST.get("universal", []) + CLAUSE_CHECKLIST.get(contract_type, [])

    results = []
    found_count = 0
    critical_missing = 0
    high_missing = 0

    for clause in clause_list:
        is_present = check_clause_presence(text[:20000], clause["keywords"])
        if is_present:
            found_count += 1

        result = {
            "id": clause["id"],
            "name": clause["name"],
            "importance": clause["importance"],
            "present": is_present,
        }
        results.append(result)

        if not is_present:
            if clause["importance"] == "critical":
                critical_missing += 1
            elif clause["importance"] == "high":
                high_missing += 1

    total = len(clause_list)
    completeness_score = round(found_count / total * 100, 1) if total > 0 else 0

    if critical_missing > 0:
        verdict = "不完整——缺少关键必备条款，建议补齐后审查"
    elif high_missing > 0:
        verdict = "基本完整——缺少部分重要条款，建议补充"
    else:
        verdict = "完整——必备条款齐全"

    return {
        "file_name": file_path.name,
        "file_path": str(file_path),
        "contract_type": contract_type,
        "total_clauses_checked": total,
        "found_clauses": found_count,
        "missing_clauses": total - found_count,
        "critical_missing": critical_missing,
        "high_missing": high_missing,
        "completeness_score": completeness_score,
        "verdict": verdict,
        "details": results,
    }


def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(
        description="检查合同条款完整性"
    )
    parser.add_argument(
        "input_file",
        help="合同文件路径，支持 PDF 或 Docx 格式"
    )
    parser.add_argument(
        "-t", "--contract-type",
        choices=["nda", "service", "procurement", "cooperation", "general", "lease", "labor"],
        help="合同类型（不指定则自动检测）"
    )
    parser.add_argument(
        "-o", "--output",
        help="输出文件路径（JSON 格式），不指定则打印到标准输出"
    )
    parser.add_argument(
        "--compact",
        action="store_true",
        help="紧凑输出模式"
    )

    args = parser.parse_args()

    try:
        result = check_completeness(args.input_file, args.contract_type)

        if args.compact:
            compact_result = {
                "contract_type": result["contract_type"],
                "completeness_score": result["completeness_score"],
                "verdict": result["verdict"],
                "critical_missing": result["critical_missing"],
                "missing": [d["name"] for d in result["details"] if not d["present"]],
            }
            result = compact_result

        json_output = json.dumps(result, ensure_ascii=False, indent=2)

        if args.output:
            with open(args.output, "w", encoding="utf-8") as f:
                f.write(json_output)
            print(f"完整性检查结果已保存到: {args.output}", file=sys.stderr)
        else:
            print(json_output)

        return 0

    except Exception as e:
        print(f"错误: {str(e)}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())

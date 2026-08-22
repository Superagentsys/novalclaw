#!/usr/bin/env python3
"""
合同元数据自动提取工具

从 PDF/Docx 格式的合同文件中自动提取关键元数据：
- 合同名称、合同类型
- 签订方信息（甲方、乙方等）
- 合同金额
- 签订日期
- 合同期限
- 合同编号
"""

import sys
import argparse
import json
import re
from pathlib import Path
from typing import Dict, List, Optional


def extract_from_pdf(file_path: str) -> str:
    """从 PDF 文件中提取文本（pdfplumber 优先，PyPDF2 备用）"""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    text_parts.append(text)
        if text_parts:
            return "\n\n".join(text_parts)
    except ImportError:
        pass
    except Exception:
        pass

    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                text_parts.append(text)
        return "\n\n".join(text_parts)
    except Exception as e:
        raise Exception(f"PDF 文件读取失败: {str(e)}")


def extract_from_docx(file_path: str) -> str:
    """从 Docx 文件中提取文本（含表格）"""
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    text_parts.append(" | ".join(row_texts))

        return "\n\n".join(text_parts)
    except Exception as e:
        raise Exception(f"Docx 文件读取失败: {str(e)}")


def extract_text(input_file: str) -> str:
    """从合同文件中提取文本内容"""
    file_path = Path(input_file)
    if not file_path.exists():
        raise FileNotFoundError(f"文件不存在: {input_file}")
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return extract_from_pdf(input_file)
    elif ext in [".docx", ".doc"]:
        return extract_from_docx(input_file)
    else:
        raise ValueError(f"不支持的文件格式: {ext}。仅支持 PDF 和 Docx 格式。")


def extract_contract_title(text: str) -> Optional[str]:
    """提取合同名称"""
    patterns = [
        r'《(.{2,60}?(?:合同|协议|书|函|备忘录|意向书|订单|标书))》',
        r'(.{2,60}?(?:合同|协议|书|函|备忘录|意向书))',
    ]
    for pattern in patterns:
        m = re.search(pattern, text[:2000])
        if m:
            title = m.group(1).strip()
            if len(title) >= 4:
                return title
    lines = text.strip().split('\n')
    for line in lines[:20]:
        line = line.strip()
        if line and len(line) >= 4 and len(line) <= 80:
            if any(k in line for k in ['合同', '协议']):
                return line
    return None


def extract_contract_type(text: str, title: Optional[str] = None) -> str:
    """识别合同类型"""
    search_text = (title or "") + " " + text[:5000]
    search_text = search_text.lower()

    type_keywords = {
        "nda": ["保密协议", "nda", "non-disclosure", "保密", "机密信息"],
        "service": ["服务合同", "服务协议", "技术服务", "咨询服务", "维护服务", "外包", "开发合同"],
        "procurement": ["采购合同", "采购协议", "供货合同", "买卖合同", "购销", "订单", "采购"],
        "cooperation": ["合作合同", "合作协议", "战略合作", "合资", "联合", "合伙", "加盟"],
        "lease": ["租赁合同", "租赁协议", "出租", "承租", "房产租赁"],
        "labor": ["劳动合同", "劳动协议", "聘用", "劳务", "雇佣"],
        "loan": ["借款合同", "贷款", "借款协议", "融资"],
        "license": ["许可合同", "授权协议", "许可协议", "商标许可", "专利许可"],
        "agency": ["代理合同", "代理协议", "委托代理", "经销"],
        "guarantee": ["担保合同", "保证合同", "担保协议", "保证"],
    }

    for ctype, keywords in type_keywords.items():
        for kw in keywords:
            if kw.lower() in search_text:
                return ctype

    return "other"


def extract_parties(text: str) -> List[Dict[str, str]]:
    """提取签订方信息"""
    parties = []

    party_patterns = [
        (r'甲方[：:]\s*(.{2,60}?)(?:[,\，\s]{1,5}(?:地址|住所|法定|统一|联系人|电话|邮箱|开户|乙方|丙方|$))', "甲方"),
        (r'乙方[：:]\s*(.{2,60}?)(?:[,\，\s]{1,5}(?:地址|住所|法定|统一|联系人|电话|邮箱|开户|甲方|丙方|$))', "乙方"),
        (r'丙方[：:]\s*(.{2,60}?)(?:[,\，\s]{1,5}(?:地址|住所|法定|统一|联系人|电话|邮箱|开户|$))', "丙方"),
        (r'买方[：:]\s*(.{2,60}?)(?:[,\，\s]{1,5}(?:地址|住所|法定|统一|$))', "买方"),
        (r'卖方[：:]\s*(.{2,60}?)(?:[,\，\s]{1,5}(?:地址|住所|法定|统一|$))', "卖方"),
        (r'供方[：:]\s*(.{2,60}?)(?:[,\，\s]{1,5}(?:地址|住所|法定|$))', "供方"),
        (r'需方[：:]\s*(.{2,60}?)(?:[,\，\s]{1,5}(?:地址|住所|法定|$))', "需方"),
    ]

    seen_parties = set()

    for text_part in [text[:3000], text]:
        for pattern, role in party_patterns:
            m = re.search(pattern, text_part)
            if m:
                name = m.group(1).strip()
                if name and name not in seen_parties and len(name) >= 2:
                    seen_parties.add(name)
                    parties.append({"role": role, "name": name})

    for match in re.finditer(r'(?:本(?:合同|协议))(?:由|的).{0,20}?(.{4,40}?)[和与及、,，]\s*(.{4,40}?)(?:共同)?签', text[:2000]):
        parties.append({"role": "签订方A", "name": match.group(1).strip()})
        parties.append({"role": "签订方B", "name": match.group(2).strip()})
        break

    return parties


def extract_contract_amount(text: str) -> Optional[Dict[str, any]]:
    """提取合同金额"""
    amount_patterns = [
        (r'(?:合同|项目|订单|协议|采购|成交|总)金额[：:]*\s*(?:人民币|RMB|￥)?\s*([\d,，.]+)\s*(万|万元|亿|元|块)?', "total"),
        (r'(?:合同|项目|订单)总价[：:]*\s*(?:人民币|RMB|￥)?\s*([\d,，.]+)\s*(万|万元|亿|元|块)?', "total"),
        (r'(?:人民币|RMB|CNY)\s*[：:]?\s*[￥¥]?\s*([\d,，.]+)\s*(万|亿|元|块)?', "total"),
        (r'[￥¥]\s*([\d,，.]+)\s*(万|亿元|万元|元|块)?', "total"),
        (r'金额[：:]*\s*(?:人民币|RMB|￥)?\s*([\d,，.]+)\s*(万|万元|亿|元|块)?', "total"),
    ]

    for pattern, amount_type in amount_patterns:
        m = re.search(pattern, text[:5000])
        if m:
            try:
                num_str = m.group(1).replace(',', '').replace('，', '').strip()
                amount = float(num_str)
                unit = m.group(2) or "元"

                if unit in ["万", "万元"]:
                    amount *= 10000
                elif unit == "亿":
                    amount *= 100000000

                return {
                    "amount": amount,
                    "unit": "CNY",
                    "raw_text": m.group(0).strip(),
                    "level": _get_amount_level(amount)
                }
            except ValueError:
                continue

    return None


def _get_amount_level(amount: float) -> str:
    """根据金额判断级别"""
    if amount <= 100000:
        return "standard"
    elif amount <= 1000000:
        return "enhanced"
    elif amount <= 10000000:
        return "strict"
    else:
        return "maximum"


def extract_contract_date(text: str) -> Optional[str]:
    """提取签订日期"""
    date_patterns = [
        r'签(?:订|署|约)(?:日期|时间)[：:]\s*(\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日)',
        r'(?:本(?:合同|协议)于|本(?:合同|协议)自)\s*(\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日)',
        r'(?:日期|时间)[：:]\s*(\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日)',
        r'(\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日)',
        r'(\d{4}[-/]\d{1,2}[-/]\d{1,2})',
    ]

    for pattern in date_patterns:
        match = re.search(pattern, text)
        if match:
            date_str = match.group(1)
            date_str = re.sub(r'\s+', '', date_str)
            date_str = date_str.replace('/', '-')
            return date_str

    return None


def extract_contract_term(text: str) -> Optional[Dict[str, any]]:
    """提取合同期限"""
    term_patterns = [
        r'(?:合同|协议|本)期限[：:]*\s*(.{5,60}?)(?:[，,。.\n]|$)',
        r'(?:有效期|履行期限)[：:]*\s*(.{5,60}?)(?:[，,。.\n]|$)',
        r'自\s*(\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日)\s*(?:起|至)\s*(?:到\s*)?(\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日)?',
    ]

    for pattern in term_patterns:
        m = re.search(pattern, text[:5000])
        if m:
            raw = m.group(0).strip()
            year_match = re.search(r'(\d+)\s*年', raw)
            month_match = re.search(r'(\d+)\s*个?月', raw)
            day_match = re.search(r'(\d+)\s*日', raw)

            years = int(year_match.group(1)) if year_match else 0
            months = int(month_match.group(1)) if month_match else 0
            total_months = years * 12 + months

            return {
                "raw_text": raw,
                "years": years,
                "months": months,
                "total_months": total_months if total_months > 0 else None,
                "is_perpetual": "永久" in raw or "长期" in raw
            }

    return None


def extract_contract_number(text: str) -> Optional[str]:
    """提取合同编号"""
    patterns = [
        r'(?:合同|协议|编号|合同号)[：:]\s*([A-Za-z0-9_\-／/]{6,40})',
    ]
    for pattern in patterns:
        m = re.search(pattern, text[:3000])
        if m:
            return m.group(1).strip()
    return None


def extract_governing_law(text: str) -> Optional[str]:
    """提取适用法律"""
    patterns = [
        r'(?:适用法律|管辖法律|法律适用)[：:]*\s*(.{5,80}?)(?:[。\n]|$)',
        r'(?:本(?:合同|协议)的)?(?:订立|履行|解释|效力).{0,20}适用(.{5,40}?)法律',
    ]
    for pattern in patterns:
        m = re.search(pattern, text[:8000])
        if m:
            return m.group(0).strip()
    return None


def extract_metadata(input_file: str) -> Dict[str, any]:
    """
    从合同文件中提取元数据

    参数:
        input_file: 合同文件路径（PDF 或 Docx 格式）

    返回:
        提取的元数据字典
    """
    file_path = Path(input_file)

    text = extract_text(input_file)
    full_text = text

    title = extract_contract_title(full_text)
    contract_type = extract_contract_type(full_text, title)
    parties = extract_parties(full_text)
    amount_info = extract_contract_amount(full_text)
    contract_date = extract_contract_date(full_text)
    contract_term = extract_contract_term(full_text)
    contract_number = extract_contract_number(full_text)
    governing_law = extract_governing_law(full_text)

    partner_level = "standard"
    if amount_info:
        if amount_info["amount"] > 10000000:
            partner_level = "basic"
    if parties:
        party_count = len(parties)
    else:
        party_count = 0

    metadata = {
        "title": title,
        "contract_type": contract_type,
        "contract_number": contract_number,
        "parties": parties,
        "party_count": party_count,
        "amount": amount_info,
        "signing_date": contract_date,
        "contract_term": contract_term,
        "governing_law": governing_law,
        "partner_level": partner_level,
    }

    result = {
        "file_name": file_path.name,
        "file_path": str(file_path),
        "file_type": file_path.suffix.lower(),
        "char_count": len(full_text),
        "metadata": metadata,
        "assessment": {
            "metadata_completeness": _assess_completeness(metadata),
            "suggested_review_level": _suggest_review_level(contract_type, amount_info, partner_level),
        }
    }

    return result


def _assess_completeness(metadata: Dict) -> Dict[str, bool]:
    """评估元数据完整度"""
    return {
        "title": metadata["title"] is not None,
        "parties": len(metadata["parties"]) >= 2,
        "amount": metadata["amount"] is not None,
        "signing_date": metadata["signing_date"] is not None,
        "contract_term": metadata["contract_term"] is not None,
        "contract_number": metadata["contract_number"] is not None,
    }


def _suggest_review_level(contract_type: str, amount_info: Optional[Dict], partner_level: str) -> str:
    """根据元数据建议审查级别"""
    if not amount_info:
        return "standard"
    level = amount_info.get("level", "standard")
    if level == "maximum" and contract_type in ["cooperation", "service"]:
        return "maximum"
    elif level in ["strict", "maximum"]:
        return "strict"
    elif level == "enhanced":
        return "enhanced"
    return "standard"


def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(
        description="从合同文件（PDF/Docx）中提取元数据"
    )
    parser.add_argument(
        "input_file",
        help="合同文件路径，支持 PDF 或 Docx 格式"
    )
    parser.add_argument(
        "-o", "--output",
        help="输出文件路径（JSON 格式），不指定则打印到标准输出"
    )
    parser.add_argument(
        "--compact",
        action="store_true",
        help="紧凑输出模式，仅输出核心字段"
    )

    args = parser.parse_args()

    try:
        result = extract_metadata(args.input_file)

        if args.compact:
            output = {
                "file_name": result["file_name"],
                "contract_type": result["metadata"]["contract_type"],
                "amount": result["metadata"]["amount"]["amount"] if result["metadata"]["amount"] else None,
                "signing_date": result["metadata"]["signing_date"],
                "suggested_review_level": result["assessment"]["suggested_review_level"],
            }
            result = output

        json_output = json.dumps(result, ensure_ascii=False, indent=2)

        if args.output:
            with open(args.output, "w", encoding="utf-8") as f:
                f.write(json_output)
            print(f"元数据已提取并保存到: {args.output}", file=sys.stderr)
        else:
            print(json_output)

        return 0

    except Exception as e:
        error_result = {
            "error": str(e),
            "file_name": Path(args.input_file).name if args.input_file else "unknown",
        }
        print(json.dumps(error_result, ensure_ascii=False, indent=2))
        return 1


if __name__ == "__main__":
    sys.exit(main())

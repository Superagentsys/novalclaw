#!/usr/bin/env python3
"""
合同标注工具

在 Docx 格式的合同文档中添加批注和高亮，标注风险条款和修改建议。
支持在段落末尾添加批注标记，并用高亮颜色区分风险等级。
"""

import sys
import argparse
import json
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX


def find_paragraphs_by_text(doc, search_text, fuzzy_match=True):
    """
    在文档中查找包含指定文本的段落

    参数:
        doc: Word 文档对象
        search_text: 要查找的文本
        fuzzy_match: 是否使用模糊匹配（包含匹配）

    返回:
        匹配的 (段落对象, 匹配文本) 列表
    """
    matches = []

    for para in doc.paragraphs:
        text = para.text
        if fuzzy_match:
            if search_text in text:
                matches.append((para, text))
        else:
            if search_text.strip() == text.strip():
                matches.append((para, text))

    return matches


def add_comment_element(paragraph, comment_text, comment_id, author="合同审查系统"):
    """
    使用 Word XML 在段落中添加真实批注

    Word 批注需要两部分配合：段落中的引用标记 + 独立的 comments 部分。
    此方法创建段落中的引用标记，comments 部分在保存前统一插入。

    参数:
        paragraph: 段落对象
        comment_text: 批注内容
        comment_id: 批注唯一标识（整数）
        author: 批注作者
    """
    if not paragraph.runs:
        return

    last_run = paragraph.runs[-1]
    run_element = last_run._r
    para_element = paragraph._p

    str_comment_id = str(comment_id)

    # 创建批注范围开始标记
    comment_start = OxmlElement('w:commentRangeStart')
    comment_start.set(qn('w:id'), str_comment_id)
    run_element.addprevious(comment_start)

    # 创建批注范围结束标记
    comment_end = OxmlElement('w:commentRangeEnd')
    comment_end.set(qn('w:id'), str_comment_id)
    # 在 run_element 之后插入
    parent = run_element.getparent()
    run_index = list(parent).index(run_element)
    parent.insert(run_index + 1, comment_end)

    # 创建批注引用标记
    comment_ref = OxmlElement('w:r')
    comment_ref_mark = OxmlElement('w:commentReference')
    comment_ref_mark.set(qn('w:id'), str_comment_id)
    comment_ref.append(comment_ref_mark)
    parent.insert(run_index + 1, comment_ref)

    return str_comment_id


def ensure_comments_part(doc):
    """确保文档包含批注部分"""
    # python-docx 不直接提供 comments 接口，通过 XML 操作添加
    # 检查是否已有 comments 部分
    comments_part = None
    for rel in doc.part.rels.values():
        if "comments" in rel.reltype:
            comments_part = rel.target_part
            break

    if comments_part is None:
        # 创建新的 comments 部分
        from docx.opc.part import Part
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        comments_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            '</w:comments>'
        )

        content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'
        partname = '/word/comments.xml'

        # 使用 doc.part.package 来添加part
        from lxml import etree
        comments_part = Part(
            partname, content_type, etree.fromstring(comments_xml.encode('utf-8')),
            doc.part.package
        )

        reltype = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
        doc.part.relate_to(comments_part, reltype)

    return comments_part


def add_comment_to_part(comments_part, comment_id, author, text):
    """向批注部分添加一条批注"""
    from lxml import etree

    comments_root = comments_part._element
    nsmap = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    }

    comment_el = etree.SubElement(comments_root, qn('w:comment'))
    comment_el.set(qn('w:id'), str(comment_id))
    comment_el.set(qn('w:author'), author)
    comment_el.set(qn('w:date'), '2026-01-01T00:00:00Z')

    # 添加批注文本段落
    p_el = etree.SubElement(comment_el, qn('w:p'))
    r_el = etree.SubElement(p_el, qn('w:r'))
    t_el = etree.SubElement(r_el, qn('w:t'))
    t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t_el.text = text


def add_highlight(paragraph, highlight_color):
    """
    为段落中的所有 run 添加高亮

    参数:
        paragraph: 段落对象
        highlight_color: WD_COLOR_INDEX 颜色常量
    """
    for run in paragraph.runs:
        run.font.highlight_color = highlight_color


def add_comment_marker(paragraph, comment_text):
    """
    在段落末尾添加批注标记（红色斜体小字）

    参数:
        paragraph: 段落对象
        comment_text: 批注文本
    """
    marker_run = paragraph.add_run(" [批注]")
    marker_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)  # 深红色
    marker_run.font.italic = True
    marker_run.font.size = Pt(8)  # 使用正确的 Pt 单位


def annotate_contract(input_file, output_file, annotations):
    """
    在合同文档中添加批注和高亮

    参数:
        input_file: 输入文件路径（Docx 格式）
        output_file: 输出文件路径
        annotations: 标注信息列表

    返回:
        成功标注的数量
    """
    from lxml import etree

    try:
        doc = Document(input_file)
        annotated_count = 0

        risk_colors = {
            "高风险": WD_COLOR_INDEX.YELLOW,
            "中风险": WD_COLOR_INDEX.BRIGHT_GREEN,
            "低风险": WD_COLOR_INDEX.GRAY_25
        }

        comment_id = 1

        for annotation in annotations:
            risk_level = annotation.get("risk_level", "中风险")
            original_text = annotation.get("original_text", "")
            suggestion = annotation.get("suggestion", "")
            explanation = annotation.get("explanation", "")

            if not original_text:
                print(f"警告: 标注项缺少 original_text，跳过", file=sys.stderr)
                continue

            matches = find_paragraphs_by_text(doc, original_text, fuzzy_match=True)

            if not matches:
                print(f"警告: 未找到文本 '{original_text[:50]}...'，跳过标注", file=sys.stderr)
                continue

            for para, matched_text in matches:
                # 添加高亮
                highlight_color = risk_colors.get(risk_level, WD_COLOR_INDEX.YELLOW)
                add_highlight(para, highlight_color)

                # 构建批注文本
                comment_parts = [f"[{risk_level}] {suggestion}"]
                if explanation:
                    comment_parts.append(f"({explanation})")
                comment_text = " ".join(comment_parts)

                # 在段落末尾添加批注标记
                add_comment_marker(para, comment_text)

                annotated_count += 1

        try:
            comments_part = ensure_comments_part(doc)
            comment_id = 1

            for annotation in annotations:
                risk_level = annotation.get("risk_level", "中风险")
                original_text = annotation.get("original_text", "")
                suggestion = annotation.get("suggestion", "")
                explanation = annotation.get("explanation", "")

                if not original_text:
                    continue

                matches = find_paragraphs_by_text(doc, original_text, fuzzy_match=True)
                if not matches:
                    continue

                comment_parts = [f"[{risk_level}] {suggestion}"]
                if explanation:
                    comment_parts.append(f"({explanation})")
                comment_text = " ".join(comment_parts)

                for para, matched_text in matches:
                    add_comment_element(para, comment_text, comment_id)
                    add_comment_to_part(comments_part, comment_id, "合同审查系统", comment_text)
                    comment_id += 1

        except Exception as e:
            print(f"注意: Word原生批注创建失败 ({str(e)})，将继续使用文内标记模式", file=sys.stderr)

        doc.save(output_file)
        return annotated_count

    except Exception as e:
        raise Exception(f"文档标注失败: {str(e)}")


def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(
        description="在合同文档中添加批注和高亮标注风险条款"
    )
    parser.add_argument(
        "input_file",
        help="输入合同文件路径（Docx 格式）"
    )
    parser.add_argument(
        "output_file",
        help="输出合同文件路径（Docx 格式）"
    )
    parser.add_argument(
        "annotations_file",
        help="标注信息文件（JSON 格式）"
    )

    args = parser.parse_args()

    try:
        with open(args.annotations_file, 'r', encoding='utf-8') as f:
            annotations = json.load(f)

        if not isinstance(annotations, list):
            raise ValueError("标注信息必须是列表格式")
        if len(annotations) == 0:
            raise ValueError("标注信息列表不能为空")

        count = annotate_contract(
            args.input_file,
            args.output_file,
            annotations
        )

        print(f"标注完成，共成功标注 {count} 个风险条款", file=sys.stderr)
        print(f"输出文件: {args.output_file}", file=sys.stderr)
        return 0

    except Exception as e:
        print(f"错误: {str(e)}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())
